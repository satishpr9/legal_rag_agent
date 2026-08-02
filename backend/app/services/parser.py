import os
from typing import List, Dict
from llama_index.core.schema import Document as LIDocument
from llama_parse import LlamaParse
from docx import Document as DocxDocument
from app.core.config import settings

class DocumentParser:
    def __init__(self):
        self.llama_parser = LlamaParse(
            api_key=settings.LLAMA_CLOUD_API_KEY,
            result_type="markdown",
            premium_mode=settings.LLAMA_PARSE_PREMIUM_MODE,
            system_prompt="This is a legal document. Carefully preserve section numbering, tables, schedules, two-column layouts, and Roman numeral sections. Ensure exact text fidelity."
        )

    def parse_pdf_pages(self, file_path: str) -> List[LIDocument]:
        parsed_docs = self.llama_parser.load_data(file_path)
        filename = os.path.basename(file_path)
        docs = []
        for i, doc in enumerate(parsed_docs):
            page_num = doc.metadata.get("page_number", i + 1) if doc.metadata else i + 1
            docs.append(LIDocument(
                text=doc.text,
                metadata={
                    "page_number": page_num,
                    "source_file": filename,
                    "file_type": "pdf"
                }
            ))
        return docs

    def parse_pdf(self, file_path: str) -> str:
        docs = self.parse_pdf_pages(file_path)
        return "\n\n".join([doc.text for doc in docs])

    def parse_docx(self, file_path: str) -> List[LIDocument]:
        docx_doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in docx_doc.paragraphs])
        filename = os.path.basename(file_path)
        return [LIDocument(
            text=text,
            metadata={
                "page_number": 1,
                "source_file": filename,
                "file_type": "docx"
            }
        )]

    def parse_txt(self, file_path: str) -> List[LIDocument]:
        filename = os.path.basename(file_path)
        ext = os.path.splitext(filename)[1].lower()[1:]
        file_type = ext if ext else "txt"
        
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
            
        return [LIDocument(
            text=text,
            metadata={
                "page_number": 1,
                "source_file": filename,
                "file_type": file_type
            }
        )]

    def parse_file_to_documents(self, file_path: str) -> List[LIDocument]:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return self.parse_pdf_pages(file_path)
        elif ext == ".docx":
            return self.parse_docx(file_path)
        elif ext in [".txt", ".md"]:
            return self.parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def parse_file(self, file_path: str) -> str:
        docs = self.parse_file_to_documents(file_path)
        return "\n\n".join([doc.text for doc in docs])

    def parse_file_pages(self, file_path: str) -> List[Dict]:
        docs = self.parse_file_to_documents(file_path)
        return [{"page_number": doc.metadata.get("page_number", 1), "text": doc.text} for doc in docs]
